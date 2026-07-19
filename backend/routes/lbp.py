"""Generator Laporan Barang Pengguna (LBP) .docx per satker — W8.

Menghasilkan dokumen LBP lengkap yang strukturnya mengikuti format resmi
(dipelajari dari LBP OIKN Tahun 2025 Audited): sampul, kata pengantar,
daftar isi, Bab I Overview (gambaran umum, dasar hukum, ruang lingkup,
kebijakan umum, kebijakan akuntansi signifikan, nilai BMN), Bab II Laporan
Barang Pengguna (posisi neraca, persediaan per akun, intra/ekstra/gabungan
per golongan, KDP, ATB, penyusutan), Bab III CaLBMN (ringkasan mutasi per
golongan Gabungan/Intra/Ekstra, BMN per akun neraca, akumulasi penyusutan,
perbandingan Laporan Barang vs Laporan Keuangan, informasi BMN lainnya),
lampiran ringkas, penutup ber-ttd Kuasa Pengguna Barang.

Seluruh ANGKA dari mesin laporan yang sudah teruji (build_dbkp_rows,
build_lbkp_rows, rekap_penyusutan, nilai persediaan FIFO) dengan lingkup
per satker (M-SCOPE). Dokumen .docx agar mudah disunting satker (narasi
kondisi lokal, tindak lanjut pemeriksaan dll.) sebelum ditandatangani.
"""
import io

from fastapi import APIRouter, Depends, Response

from auth_utils import require_user
from db import db
from lbp_utils import (DASAR_HUKUM_LBP, UPAYA_PERBAIKAN_LBP,
                       baris_akumulasi_per_akun, baris_mutasi_golongan,
                       baris_perbandingan_lb_lk, baris_posisi_per_akun,
                       fmt_rp, kebijakan_akuntansi_lbp, label_periode_lbp,
                       rupiah_terbilang, struktur_daftar_isi_lengkap,
                       susun_mutasi_per_transaksi, tanggal_akhir_periode)
from shared_utils import (ambang_kapitalisasi, filter_aset_perhitungan,
                          kode_satker_user,
                          pengaturan_kop, resolve_penandatangan_kpb,
                          scope_query_aset, scope_query_field_satker)

lbp_router = APIRouter()

_BIRU = "1F4E79"
_ABU = "D9E2F3"


def _set_shading(cell, hexcolor):
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:fill"), hexcolor)
    cell._tc.get_or_add_tcPr().append(shd)


def _doc_baru():
    from docx import Document
    from docx.shared import Cm, Pt
    d = Document()
    st = d.styles["Normal"]
    st.font.name = "Cambria"
    st.font.size = Pt(11)
    for s in d.sections:
        s.top_margin = Cm(2.5)
        s.bottom_margin = Cm(2.5)
        s.left_margin = Cm(2.5)
        s.right_margin = Cm(2.5)
    return d


def _footer_halaman(d):
    """Nomor halaman otomatis di footer tengah (field PAGE)."""
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    for s in d.sections:
        p = s.footer.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        fld = OxmlElement("w:fldSimple")
        fld.set(qn("w:instr"), "PAGE")
        p._p.append(fld)


def _h(d, teks, level=1):
    from docx.shared import Pt, RGBColor
    p = d.add_heading(teks, level=level)
    for run in p.runs:
        run.font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)
        run.font.name = "Cambria"
        run.font.size = Pt(15 if level == 1 else 12)
    return p


def _p(d, teks, bold=False, center=False, size=None, space_after=6):
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Pt
    p = d.add_paragraph()
    run = p.add_run(teks)
    run.bold = bold
    if size:
        run.font.size = Pt(size)
    if center:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(space_after)
    return p


def _daftar(d, items):
    for it in items:
        d.add_paragraph(it, style="List Number")


def _tabel(d, header, rows, lebar=None, align_kanan=None):
    """Tabel bergaris: header biru putih tebal + isi; align_kanan = set
    indeks kolom rata kanan (angka)."""
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Cm, Pt, RGBColor
    t = d.add_table(rows=1, cols=len(header))
    t.style = "Table Grid"
    kanan = align_kanan or set()
    for i, judul in enumerate(header):
        c = t.rows[0].cells[i]
        c.text = ""
        run = c.paragraphs[0].add_run(str(judul))
        run.bold = True
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        c.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        _set_shading(c, _BIRU)
    for ridx, row in enumerate(rows):
        cells = t.add_row().cells
        for i, val in enumerate(row):
            c = cells[i]
            c.text = ""
            run = c.paragraphs[0].add_run(str(val))
            run.font.size = Pt(9)
            if i in kanan:
                c.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
            if ridx % 2 == 1:
                _set_shading(c, _ABU)
    if lebar:
        for i, w in enumerate(lebar):
            for r in t.rows:
                r.cells[i].width = Cm(w)
    d.add_paragraph()
    return t


def _tabel_lbkp(d, judul, rows, total):
    """Satu seksi laporan per golongan (saldo awal/tambah/kurang/akhir)."""
    _p(d, judul, bold=True)
    header = ["Gol", "Uraian", "Jml Awal", "Nilai Awal", "Jml Tambah",
              "Nilai Tambah", "Jml Kurang", "Nilai Kurang", "Jml Akhir",
              "Nilai Akhir"]
    baris = [[r.get("golongan", ""), r.get("uraian", ""),
              fmt_rp(r.get("jumlah_awal")), fmt_rp(r.get("nilai_awal")),
              fmt_rp(r.get("jumlah_tambah")), fmt_rp(r.get("nilai_tambah")),
              fmt_rp(r.get("jumlah_kurang")), fmt_rp(r.get("nilai_kurang")),
              fmt_rp(r.get("jumlah_akhir")), fmt_rp(r.get("nilai_akhir"))]
             for r in rows]
    baris.append(["", "TOTAL",
                  fmt_rp(total.get("jumlah_awal")), fmt_rp(total.get("nilai_awal")),
                  fmt_rp(total.get("jumlah_tambah")), fmt_rp(total.get("nilai_tambah")),
                  fmt_rp(total.get("jumlah_kurang")), fmt_rp(total.get("nilai_kurang")),
                  fmt_rp(total.get("jumlah_akhir")), fmt_rp(total.get("nilai_akhir"))])
    _tabel(d, header, baris, align_kanan=set(range(2, 10)))


@lbp_router.get("/pelaporan/lbp-docx")
async def generate_lbp_docx(tahun: int, semester: int = 0,
                            user: dict = Depends(require_user)):
    """Rakit LBP .docx per satker untuk periode tahun/semester."""
    from kodefikasi_utils import GOLONGAN_DEFAULTS
    from pemeliharaan_utils import rentang_periode
    from pembukuan_utils import (build_dbkp_rows, build_lbkp_rows,
                                 parse_harga, tombstones_penghapusan)
    from penilaian_utils import MASA_MANFAAT_DEFAULT, rekap_penyusutan
    from persediaan_akun_utils import akun_persediaan
    from persediaan_utils import (nilai_persediaan_dari_batches,
                                  stok_dari_batches)
    from routes.reports import _fmt_tanggal_id

    sem = semester if semester in (1, 2) else None
    dari, sampai, label_rentang = rentang_periode(tahun, sem)
    per_iso = tanggal_akhir_periode(tahun, sem)
    label = label_periode_lbp(tahun, sem)

    kode = kode_satker_user(user)
    settings = await pengaturan_kop(kode_satker=kode)
    ttd = await resolve_penandatangan_kpb(settings, per_iso=per_iso)
    ambang = await ambang_kapitalisasi()

    # ── Data aset (scoped satker) ──
    q = await filter_aset_perhitungan(await scope_query_aset(user, {}))
    assets = await db.assets.find(q, {
        "_id": 0, "id": 1, "asset_code": 1, "NUP": 1, "asset_name": 1,
        "purchase_price": 1, "purchase_date": 1, "created_at": 1,
        "dihapus": 1, "penghapusan": 1, "nilai_wajar_terakhir": 1,
        "revaluasi": 1, "location": 1}).to_list(500000)
    aktif = [a for a in assets if not a.get("dihapus")]

    uraian_map = {k: u for k, u in GOLONGAN_DEFAULTS}
    async for k in db.kodefikasi.find({"level": 1},
                                      {"_id": 0, "kode": 1, "uraian": 1}):
        if k.get("uraian"):
            uraian_map[k["kode"]] = k["uraian"]
    peta_akun = {}
    async for m in db.akun_bas.find({}, {"_id": 0}):
        peta_akun[str(m.get("golongan"))] = m

    # Posisi (DBKP) + persediaan per akun 117xxx
    rows_dbkp, total_dbkp = build_dbkp_rows(aktif, uraian_map, ambang)
    peta_psd_override = {}
    async for m in db.persediaan_akun.find({}, {"_id": 0}):
        peta_psd_override[m.get("sub_kelompok", "")] = m
    psd_akun = {}
    psd_jumlah = 0
    async for it in db.persediaan.find(
            scope_query_field_satker(user),
            {"_id": 0, "kode_barang": 1, "batches": 1}):
        info = akun_persediaan(it.get("kode_barang"), peta_psd_override)
        r = psd_akun.setdefault(info["akun"], {
            "akun": info["akun"], "uraian": info["uraian"], "nilai": 0.0})
        r["nilai"] += nilai_persediaan_dari_batches(it.get("batches"))
        psd_jumlah += stok_dari_batches(it.get("batches"))
    persediaan_per_akun = sorted(psd_akun.values(), key=lambda x: x["akun"])
    psd_nilai = sum(x["nilai"] for x in persediaan_per_akun)

    # Mutasi periode (LBKP) — tombstone audit + SK penghapusan
    tombstones = []
    async for t in db.audit_logs.find(
            {"action": "delete"},
            {"_id": 0, "asset_code": 1, "timestamp": 1, "changes": 1}):
        nilai = 0.0
        for c in t.get("changes") or []:
            if c.get("field") == "purchase_price":
                nilai = parse_harga(c.get("from"))
                break
        tombstones.append({"asset_code": t.get("asset_code"),
                           "timestamp": t.get("timestamp"), "nilai": nilai})
    tombstones += tombstones_penghapusan(assets, ambang=ambang)
    per_kelas, _tanpa_nilai = build_lbkp_rows(
        assets, tombstones, dari, sampai, uraian_map, ambang)

    # Penyusutan per golongan
    peta_masa = dict(MASA_MANFAAT_DEFAULT)
    async for m in db.masa_manfaat.find({}, {"_id": 0}):
        peta_masa[m["kode"]] = int(m["tahun"])
    susut = rekap_penyusutan(aktif, per_iso, peta=peta_masa,
                             uraian_golongan=uraian_map)

    # Susunan tabel CaLBMN
    posisi_akun = baris_posisi_per_akun(rows_dbkp, persediaan_per_akun,
                                        peta_akun)
    akumulasi = baris_akumulasi_per_akun(susut)
    banding = baris_perbandingan_lb_lk(posisi_akun, akumulasi)
    mutasi_gol = baris_mutasi_golongan(per_kelas)
    netto = posisi_akun["total"] - akumulasi["total"]

    # Jurnal Buku Barang periode (per kode transaksi, utk CaLBMN)
    ids_scoped = [a.get("id") for a in assets if a.get("id")]
    jurnal_periode = await db.mutasi_bmn.find(
        {"asset_id": {"$in": ids_scoped},
         "tanggal_buku": {"$gte": dari, "$lte": sampai}},
        {"_id": 0, "kode_barang": 1, "kode_transaksi": 1, "jumlah": 1,
         "nilai": 1}).to_list(100000)
    jurnal_per_gol = {}
    for j in jurnal_periode:
        g = str(j.get("kode_barang") or "")[:1] or "?"
        jurnal_per_gol.setdefault(g, []).append(j)

    # Register pendukung "Informasi BMN Lainnya"
    f_satker = scope_query_field_satker(user)
    n_sengketa = await db.pengamanan_kasus.count_documents(
        {**f_satker, "status": {"$ne": "selesai"}})
    n_idle = await db.bmn_idle.count_documents(
        {**f_satker, "status": {"$in": ["klarifikasi", "usul_serah"]}})
    n_psp = await db.psp.count_documents(f_satker)
    n_tertib = await db.penertiban.count_documents(
        {**f_satker, "status": "berjalan"})

    # Data LAMPIRAN
    sengketa_rows = await db.pengamanan_kasus.find(
        {**f_satker, "status": {"$ne": "selesai"}},
        {"_id": 0, "asset_name": 1, "asset_code": 1, "kategori": 1,
         "status": 1, "pihak_lawan": 1}).to_list(50)
    polis_rows = await db.pengamanan_polis.find(
        {**f_satker}, {"_id": 0, "nomor_polis": 1, "penanggung": 1,
                       "asset_name": 1, "nilai_pertanggungan": 1,
                       "premi": 1, "mulai": 1, "berakhir": 1}).to_list(100)
    pnbp_rows = []
    total_pnbp = 0.0
    async for pm in db.pemanfaatan.find(
            {**f_satker}, {"_id": 0, "bentuk": 1, "mitra": 1,
                           "kontribusi": 1}):
        for kt in pm.get("kontribusi") or []:
            tgl = str(kt.get("tanggal") or "")[:10]
            if tgl and not (dari <= tgl <= sampai):
                continue
            nilai_k = float(kt.get("jumlah") or 0)
            total_pnbp += nilai_k
            if len(pnbp_rows) < 100:
                pnbp_rows.append([tgl, pm.get("bentuk") or "",
                                  pm.get("mitra") or "",
                                  kt.get("ntpn") or "", fmt_rp(nilai_k)])
    hilang_rows = await db.usulan_penghapusan.find(
        {**f_satker, "jalur": "tidak_ditemukan"},
        {"_id": 0, "asset_code": 1, "NUP": 1, "asset_name": 1,
         "status": 1, "nomor_sk": 1}).to_list(100)
    rumah_negara = [a for a in aktif
                    if "rumah negara" in str(a.get("asset_name") or "").lower()][:100]
    transfer_rows = [j for j in jurnal_periode
                     if str(j.get("kode_transaksi")) in ("102", "302")]
    hibah_rows = [j for j in jurnal_periode
                  if str(j.get("kode_transaksi")) == "103"]

    nama_satker = (settings.get("nama_sub_unit")
                   or settings.get("nama_unit_organisasi")
                   or settings.get("nama_instansi") or "Satuan Kerja")
    instansi = settings.get("nama_instansi") or nama_satker

    # ══ Rakit dokumen ══
    d = _doc_baru()

    # SAMPUL
    for _ in range(6):
        d.add_paragraph()
    _p(d, "LAPORAN BARANG PENGGUNA", bold=True, center=True, size=24,
       space_after=2)
    _p(d, f"{label.upper()}", bold=True, center=True, size=16)
    d.add_paragraph()
    _p(d, instansi.upper(), bold=True, center=True, size=14, space_after=2)
    _p(d, nama_satker.upper(), center=True, size=12, space_after=2)
    if kode:
        _p(d, f"Kode Satker: {kode}", center=True, size=11)
    _p(d, f"Posisi per {_fmt_tanggal_id(per_iso)}", center=True, size=11)
    d.add_page_break()

    # KATA PENGANTAR
    _h(d, "KATA PENGANTAR", 1)
    _p(d, "Sebagaimana diamanatkan Undang-Undang Nomor 17 Tahun 2003 "
          "tentang Keuangan Negara dan Undang-Undang Nomor 1 Tahun 2004 "
          "tentang Perbendaharaan Negara, Pengguna/Kuasa Pengguna Barang "
          "mempunyai tugas mengelola barang milik negara serta menyusun "
          "dan menyampaikan laporannya secara periodik.")
    _p(d, f"Penyusunan Laporan Barang Pengguna (LBP) {nama_satker} {label} "
          "mengacu pada PP Nomor 27 Tahun 2014 jo. PP Nomor 28 Tahun 2020, "
          "PP Nomor 71 Tahun 2010 tentang Standar Akuntansi Pemerintahan, "
          "PMK Nomor 181/PMK.06/2016 tentang Penatausahaan BMN, dan PMK "
          "Nomor 118/PMK.06/2018 tentang Tata Cara Rekonsiliasi BMN.")
    _p(d, "Upaya perbaikan kualitas LBP yang dilakukan antara lain:")
    _daftar(d, UPAYA_PERBAIKAN_LBP)
    _p(d, "Semoga Laporan Barang Pengguna ini memberikan informasi yang "
          "bermanfaat bagi para pemangku kepentingan serta manajemen dalam "
          "proses pengambilan keputusan pengelolaan BMN.")
    d.add_paragraph()
    _p(d, f"{settings.get('tempat_laporan') or ''} "
          f"{_fmt_tanggal_id(per_iso)}".strip(), center=False)
    _p(d, "Kuasa Pengguna Barang,", space_after=40)
    _p(d, ttd.get("nama") or "………………………………", bold=True, space_after=0)
    if ttd.get("nip"):
        _p(d, f"NIP. {ttd['nip']}", space_after=0)
    d.add_page_break()

    _footer_halaman(d)

    # DAFTAR ISI (outline lengkap)
    _h(d, "DAFTAR ISI", 1)
    for level, judul in struktur_daftar_isi_lengkap():
        _p(d, ("    " * (level - 1)) + judul, bold=level == 1, space_after=2)
    d.add_page_break()

    # SURAT PENGANTAR
    _h(d, "SURAT PENGANTAR", 1)
    _p(d, f"Bersama ini disampaikan Laporan Barang Pengguna {nama_satker} "
          f"{label}, dengan posisi per {_fmt_tanggal_id(per_iso)}, sebagai "
          "bahan penyusunan Laporan Barang Pengguna tingkat Kementerian/"
          "Lembaga dan bahan penyusunan neraca, untuk dipergunakan "
          "sebagaimana mestinya.")
    _p(d, "Laporan ini disusun dari data penatausahaan BMN aplikasi AMAN "
          "yang telah melalui rekonsiliasi internal dengan unit akuntansi "
          "keuangan.")
    d.add_paragraph()
    _p(d, "Kuasa Pengguna Barang,", space_after=40)
    _p(d, ttd.get("nama") or "………………………………", bold=True, space_after=0)
    if ttd.get("nip"):
        _p(d, f"NIP. {ttd['nip']}", space_after=0)
    d.add_page_break()

    # BAB I — OVERVIEW
    _h(d, "I. OVERVIEW LAPORAN BARANG PENGGUNA", 1)
    _h(d, "1.1 Gambaran Umum", 2)
    _p(d, f"LBP {nama_satker} {label} disusun berdasarkan data BMN yang "
          "ditatausahakan dalam aplikasi AMAN dan telah melalui proses "
          "rekonsiliasi internal dengan unit akuntansi keuangan. Nilai BMN "
          f"pada laporan posisi BMN di neraca per {_fmt_tanggal_id(per_iso)} "
          f"(netto setelah penyusutan) adalah sebesar "
          f"{rupiah_terbilang(netto)}.")
    _h(d, "1.2 Dasar Hukum", 2)
    _daftar(d, DASAR_HUKUM_LBP)
    _h(d, "1.3 Ruang Lingkup Laporan", 2)
    _p(d, "LBP ini menyajikan nilai BMN secara komprehensif, baik "
          "intrakomptabel (memenuhi syarat kapitalisasi dan disajikan di "
          "neraca) maupun ekstrakomptabel (tidak memenuhi syarat "
          "kapitalisasi), berdasarkan klasifikasi Bagan Akun Standar dan "
          "penggolongan barang KMK 29/PMK.6/2010 beserta perubahannya.")
    _h(d, "1.4 Kebijakan Umum Penatausahaan BMN", 2)
    _p(d, "BMN adalah semua barang yang dibeli atau diperoleh atas beban "
          "APBN atau berasal dari perolehan lainnya yang sah. BMN dicatat "
          "dan dilaporkan sesuai asas fungsional, kepastian hukum, "
          "transparansi, efisiensi, akuntabilitas, dan kepastian nilai.")
    _h(d, "1.5 Kebijakan Akuntansi yang Signifikan", 2)
    for sub in kebijakan_akuntansi_lbp(ambang):
        _p(d, sub["judul"], bold=True, space_after=2)
        for teks in sub["isi"]:
            _p(d, teks)
        if sub["daftar"]:
            for butir in sub["daftar"]:
                d.add_paragraph(butir, style="List Bullet")
    _p(d, "Jumlah Satuan Kerja", bold=True, space_after=2)
    _p(d, (f"Laporan ini mencakup 1 (satu) satuan kerja dengan kode "
           f"satker {kode}." if kode else
           "Laporan ini disusun pada lingkup seluruh satuan kerja yang "
           "tercatat pada aplikasi (akun lintas-satker)."))
    _h(d, "1.6 Nilai Barang Milik Negara", 2)
    total_intra = float(total_dbkp.get("nilai_intra") or 0) + psd_nilai
    total_ekstra = float(total_dbkp.get("nilai_ekstra") or 0)
    _tabel(d, ["No", "Uraian", "Nilai (Rp)"], [
        ["1", "BMN di Neraca — Intrakomptabel (bruto, termasuk persediaan)",
         fmt_rp(total_intra)],
        ["2", "BMN Ekstrakomptabel (bruto)", fmt_rp(total_ekstra)],
        ["3", "Akumulasi Penyusutan/Amortisasi", fmt_rp(-akumulasi["total"])],
        ["", "Total (netto)", fmt_rp(total_intra + total_ekstra
                                     - akumulasi["total"])],
    ], align_kanan={2})

    # Perkembangan nilai BMN vs saldo awal periode (bruto, tanpa persediaan)
    gab_rows, gab_total = per_kelas.get("gabungan") or ([], {})
    nilai_awal_gab = float(gab_total.get("nilai_awal") or 0)
    nilai_akhir_gab = float(gab_total.get("nilai_akhir") or 0)
    _p(d, "Perkembangan Nilai BMN", bold=True, space_after=2)
    _p(d, f"Dibandingkan saldo awal periode, nilai BMN (bruto, di luar "
          f"persediaan) mengalami "
          f"{'peningkatan' if nilai_akhir_gab >= nilai_awal_gab else 'penurunan'} "
          f"sebesar {rupiah_terbilang(abs(nilai_akhir_gab - nilai_awal_gab))}:")
    _tabel(d, ["No", "Uraian", "Saldo Awal Periode", "Saldo Akhir Periode",
               "Peningkatan/(Penurunan)"], [
        ["1", "BMN Gabungan (Intra + Ekstra, bruto)",
         fmt_rp(nilai_awal_gab), fmt_rp(nilai_akhir_gab),
         fmt_rp(nilai_akhir_gab - nilai_awal_gab)],
    ], align_kanan={2, 3, 4})

    # Komposisi BMN per jenis aset (narasi ringkas per golongan)
    _p(d, "Komposisi BMN per Jenis Aset", bold=True, space_after=2)
    if psd_nilai:
        _p(d, f"Persediaan per {_fmt_tanggal_id(per_iso)} sebesar "
              f"{rupiah_terbilang(psd_nilai)}.")
    for blok in mutasi_gol:
        akhir_n = blok["baris"][3][2]
        _p(d, f"{blok['uraian']} (golongan {blok['golongan']}) — saldo "
              f"akhir gabungan sebesar {rupiah_terbilang(akhir_n)}.")
    d.add_page_break()

    # BAB II — LAPORAN BARANG PENGGUNA
    _h(d, "II. LAPORAN BARANG PENGGUNA", 1)
    _h(d, "a. Laporan Posisi BMN di Neraca", 2)
    _p(d, f"Posisi per {_fmt_tanggal_id(per_iso)} — rincian per golongan "
          "barang (nilai buku bruto):")
    _tabel(d, ["Gol", "Uraian", "Jml Intra", "Nilai Intra", "Jml Ekstra",
               "Nilai Ekstra", "Jml Total", "Nilai Total"],
           [[r.get("golongan"), r.get("uraian"),
             fmt_rp(r.get("jumlah_intra")), fmt_rp(r.get("nilai_intra")),
             fmt_rp(r.get("jumlah_ekstra")), fmt_rp(r.get("nilai_ekstra")),
             fmt_rp(r.get("jumlah_total")), fmt_rp(r.get("nilai_total"))]
            for r in rows_dbkp] +
           [["", "TOTAL", fmt_rp(total_dbkp.get("jumlah_intra")),
             fmt_rp(total_dbkp.get("nilai_intra")),
             fmt_rp(total_dbkp.get("jumlah_ekstra")),
             fmt_rp(total_dbkp.get("nilai_ekstra")),
             fmt_rp(total_dbkp.get("jumlah_total")),
             fmt_rp(total_dbkp.get("nilai_total"))]],
           align_kanan=set(range(2, 8)))

    _h(d, "b. Laporan Persediaan", 2)
    if persediaan_per_akun:
        _tabel(d, ["Akun", "Uraian", "Nilai (Rp)"],
               [[x["akun"], x["uraian"], fmt_rp(x["nilai"])]
                for x in persediaan_per_akun] +
               [["", "Total", fmt_rp(psd_nilai)]], align_kanan={2})
    else:
        _p(d, "Nihil — tidak terdapat saldo persediaan pada periode ini.")

    for judul, kunci in (("c. Laporan BMN Intrakomptabel", "intra"),
                         ("d. Laporan BMN Ekstrakomptabel", "ekstra"),
                         ("e. Laporan BMN Gabungan", "gabungan")):
        _h(d, judul, 2)
        rows_k, total_k = per_kelas.get(kunci) or ([], {})
        if rows_k:
            _tabel_lbkp(d, f"Periode {label_rentang}", rows_k, total_k)
        else:
            _p(d, "Nihil pada periode ini.")

    _h(d, "f. Laporan Konstruksi Dalam Pengerjaan (KDP)", 2)
    kdp = [r for r in rows_dbkp if str(r.get("golongan")) == "7"]
    if kdp:
        _tabel(d, ["Uraian", "Jumlah", "Nilai (Rp)"],
               [[r.get("uraian"), fmt_rp(r.get("jumlah_total")),
                 fmt_rp(r.get("nilai_total"))] for r in kdp],
               align_kanan={1, 2})
    else:
        _p(d, "Nihil — tidak terdapat KDP pada periode ini.")

    _h(d, "g. Laporan Aset Tak Berwujud", 2)
    atb = [r for r in rows_dbkp if str(r.get("golongan")) == "8"]
    if atb:
        _tabel(d, ["Uraian", "Jumlah", "Nilai (Rp)"],
               [[r.get("uraian"), fmt_rp(r.get("jumlah_total")),
                 fmt_rp(r.get("nilai_total"))] for r in atb],
               align_kanan={1, 2})
    else:
        _p(d, "Nihil — tidak terdapat Aset Tak Berwujud pada periode ini.")

    for judul_n in ("h. Laporan Barang Bersejarah",
                    "i. Laporan Barang BPYDS",
                    "j. Laporan Barang Hibah DK/TP"):
        _h(d, judul_n, 2)
        _p(d, "Nihil — tidak terdapat data pada kategori ini di periode "
              "pelaporan. (Lengkapi bila satker memiliki barang pada "
              "kategori ini.)")

    _h(d, "k. Laporan Penyusutan dan Amortisasi", 2)
    per_gol_susut = (susut or {}).get("per_golongan") or []
    if per_gol_susut:
        _tabel(d, ["Gol", "Uraian", "Jumlah", "Nilai Perolehan",
                   "Akumulasi Penyusutan", "Nilai Buku"],
               [[r.get("golongan"), r.get("uraian"), fmt_rp(r.get("jumlah")),
                 fmt_rp(r.get("nilai_perolehan")), fmt_rp(r.get("akumulasi")),
                 fmt_rp(r.get("nilai_buku"))] for r in per_gol_susut] +
               [["", "TOTAL", fmt_rp(susut["total"].get("jumlah")),
                 fmt_rp(susut["total"].get("nilai_perolehan")),
                 fmt_rp(susut["total"].get("akumulasi")),
                 fmt_rp(susut["total"].get("nilai_buku"))]],
               align_kanan=set(range(2, 6)))
    else:
        _p(d, "Nihil — tidak terdapat objek penyusutan pada periode ini.")
    d.add_page_break()

    # BAB III — CaLBMN
    _h(d, "III. CATATAN ATAS LAPORAN BARANG MILIK NEGARA", 1)
    _h(d, "3.1 Pendahuluan", 2)
    _p(d, f"Catatan atas Laporan BMN ini menyajikan penjelasan pos-pos "
          f"LBP {nama_satker} {label}, meliputi ringkasan mutasi per "
          "golongan barang, penyajian per akun neraca, serta perbandingan "
          "nilai antara Laporan Barang dan Laporan Keuangan.")
    _h(d, "3.2 Ringkasan Mutasi BMN per Golongan", 2)
    if mutasi_gol:
        for blok in mutasi_gol:
            g = blok["golongan"]
            akhir_n = blok["baris"][3][2]
            _p(d, f"Golongan {g} — {blok['uraian']}", bold=True)
            _p(d, f"Saldo {blok['uraian']} pada LBP {nama_satker} {label} "
                  f"per {_fmt_tanggal_id(per_iso)} sebesar "
                  f"{rupiah_terbilang(akhir_n)}, dengan ringkasan mutasi:")
            _tabel(d, ["Uraian", "Unit Gab.", "Nilai Gabungan",
                       "Unit Intra", "Nilai Intra", "Unit Ekstra",
                       "Nilai Ekstra"],
                   [[b[0], fmt_rp(b[1]), fmt_rp(b[2]), fmt_rp(b[3]),
                     fmt_rp(b[4]), fmt_rp(b[5]), fmt_rp(b[6])]
                    for b in blok["baris"]],
                   align_kanan=set(range(1, 7)))
            # Rincian mutasi per kode transaksi (jurnal Buku Barang)
            mtx = susun_mutasi_per_transaksi(
                jurnal_per_gol.get(g, []),
                saldo_awal_qty=blok["baris"][0][1],
                saldo_awal_nilai=blok["baris"][0][2])
            if len(mtx["baris"]) > 1:
                _p(d, "Rincian mutasi yang ditatausahakan (Buku Barang):")
                _tabel(d, ["Kode", "Uraian", "Kuantitas", "Rupiah"],
                       [[k, u, fmt_rp(q), fmt_rp(n)]
                        for k, u, q, n in mtx["baris"]] +
                       [["", "Total", fmt_rp(mtx["total"][0]),
                         fmt_rp(mtx["total"][1])]],
                       align_kanan={2, 3})
            # Rincian per NUP untuk Tanah (pola dokumen contoh)
            if g == "2":
                tanah = [a for a in aktif
                         if str(a.get("asset_code") or "")[:1] == "2"]
                if tanah:
                    _p(d, "Rincian tanah yang ditatausahakan "
                          f"({len(tanah)} NUP"
                          f"{', 150 pertama ditampilkan' if len(tanah) > 150 else ''}):")
                    from pembukuan_utils import parse_harga as _ph
                    _tabel(d, ["No", "Kode Barang", "Nama Barang", "NUP",
                               "Nilai", "Keterangan"],
                           [[str(i + 1), t.get("asset_code") or "",
                             t.get("asset_name") or "", t.get("NUP") or "",
                             fmt_rp(_ph(t.get("purchase_price"))),
                             t.get("location") or ""]
                            for i, t in enumerate(tanah[:150])],
                           align_kanan={4})
    else:
        _p(d, "Tidak terdapat mutasi BMN pada periode ini.")
    _h(d, "3.3 BMN per Akun Neraca", 2)
    baris_akun = []
    for seksi, kunci in (("I. Aset Lancar", "aset_lancar"),
                         ("II. Aset Tetap", "aset_tetap"),
                         ("III. Aset Lainnya", "aset_lainnya")):
        isi = posisi_akun.get(kunci) or []
        if not isi:
            continue
        baris_akun.append(["", seksi, ""])
        for akun, uraian, nilai in isi:
            baris_akun.append([akun, uraian, fmt_rp(nilai)])
        baris_akun.append(["", f"Jumlah {seksi.split('. ', 1)[-1]}",
                           fmt_rp(posisi_akun["subtotal"][kunci])])
    baris_akun.append(["", "JUMLAH TOTAL (bruto)",
                       fmt_rp(posisi_akun["total"])])
    _tabel(d, ["Akun", "Uraian Neraca", "Nilai (Rp)"], baris_akun,
           align_kanan={2})
    if akumulasi["baris"]:
        _p(d, "Rincian akumulasi penyusutan/amortisasi per akun:")
        _tabel(d, ["Akun", "Uraian", "Nilai (Rp)"],
               [[a, u, fmt_rp(n)] for a, u, n in akumulasi["baris"]] +
               [["", "Jumlah Akumulasi", fmt_rp(akumulasi["total"])]],
               align_kanan={2})
    _p(d, f"Nilai BMN netto per {_fmt_tanggal_id(per_iso)} sebesar "
          f"{rupiah_terbilang(netto)}.")
    _h(d, "3.4 Perbandingan Laporan Barang dan Laporan Keuangan", 2)
    baris_lk = []
    for judul, rws in banding["seksi"]:
        baris_lk.append(["", judul, "", "", ""])
        for akun, uraian, lb, lk, selisih in rws:
            baris_lk.append([akun, uraian, fmt_rp(lb), fmt_rp(lk),
                             fmt_rp(selisih)])
    tlb, tlk, tsel = banding["total"]
    baris_lk.append(["", "JUMLAH", fmt_rp(tlb), fmt_rp(tlk), fmt_rp(tsel)])
    _tabel(d, ["Akun", "Uraian", "Laporan Barang", "Laporan Keuangan",
               "Selisih"], baris_lk, align_kanan={2, 3, 4})
    _p(d, "Berdasarkan rekapitulasi di atas, tidak terdapat perbedaan "
          "nilai antara Laporan Barang dan Laporan Keuangan (satu basis "
          "data penatausahaan). Bila terdapat selisih pada rekonsiliasi "
          "eksternal (SAKTI/e-Rekon), uraikan penjelasannya di sini.")
    _h(d, "3.5 Informasi BMN Lainnya", 2)
    _p(d, "a. Ringkasan register pengelolaan", bold=True, space_after=2)
    _tabel(d, ["No", "Informasi", "Jumlah"], [
        ["1", "SK Penetapan Status Penggunaan (PSP) tercatat", fmt_rp(n_psp)],
        ["2", "BMN idle dalam proses klarifikasi/usul serah", fmt_rp(n_idle)],
        ["3", "Kasus pengamanan/sengketa belum selesai", fmt_rp(n_sengketa)],
        ["4", "Penertiban Wasdal berjalan", fmt_rp(n_tertib)],
    ], align_kanan={2})
    _p(d, "b. Dokumen Sumber Tanah", bold=True, space_after=2)
    tanah_semua = [a for a in aktif
                   if str(a.get("asset_code") or "")[:1] == "2"]
    if tanah_semua:
        from pembukuan_utils import parse_harga as _ph2
        _tabel(d, ["No", "Uraian", "Jumlah NUP", "Nilai (Rp)"], [
            ["1", "Tanah Barang Milik Negara (tercatat aplikasi)",
             fmt_rp(len(tanah_semua)),
             fmt_rp(sum(_ph2(t.get("purchase_price")) for t in tanah_semua))],
        ], align_kanan={2, 3})
        _p(d, "Rincian dokumen sumber (sertipikat/HPL/BAST perolehan) "
              "dilengkapi satker sesuai dokumen kepemilikan masing-masing "
              "bidang.")
    else:
        _p(d, "Nihil — tidak terdapat BMN berupa tanah.")
    _p(d, "c. BMN Bersengketa", bold=True, space_after=2)
    if sengketa_rows:
        _tabel(d, ["No", "Aset", "Kategori", "Pihak Lawan", "Status"],
               [[str(i + 1),
                 f"{s.get('asset_name') or ''} ({s.get('asset_code') or ''})",
                 s.get("kategori") or "", s.get("pihak_lawan") or "",
                 s.get("status") or ""]
                for i, s in enumerate(sengketa_rows)])
    else:
        _p(d, "Nihil — tidak terdapat BMN bersengketa pada periode ini.")
    _p(d, "d. Permasalahan Penatausahaan BMN", bold=True, space_after=2)
    _p(d, "(Uraikan permasalahan penatausahaan yang dihadapi satker — "
          "mis. aset belum ber-PSP, dokumen kepemilikan dalam proses, "
          "selisih hasil inventarisasi — lengkapi sebelum tanda tangan.)")
    _p(d, "e. Langkah-langkah Strategis Penyelesaian Masalah", bold=True,
       space_after=2)
    _p(d, "(Uraikan rencana aksi/langkah strategis penyelesaian atas "
          "permasalahan pada huruf d.)")

    _h(d, "3.6 Tindak Lanjut Temuan Pemeriksaan", 2)
    _p(d, "Rekapitulasi temuan pemeriksaan (BPK/APIP) atas laporan "
          "keuangan/BMN dan tindak lanjutnya — lengkapi sesuai LHP:")
    _tabel(d, ["No", "Temuan", "Rekomendasi", "Tindak Lanjut", "Status"],
           [["1", "…", "…", "…", "…"]])
    d.add_page_break()

    # LAMPIRAN
    _h(d, "LAMPIRAN", 1)
    _h(d, "a. Laporan PNBP dari Pengelolaan BMN", 2)
    if pnbp_rows:
        _p(d, f"Total penerimaan periode {label_rentang}: "
              f"{rupiah_terbilang(total_pnbp)}.")
        _tabel(d, ["Tanggal", "Bentuk", "Mitra", "NTPN", "Nilai (Rp)"],
               pnbp_rows, align_kanan={4})
    else:
        _p(d, "Nihil — tidak terdapat setoran PNBP dari pemanfaatan BMN "
              "pada periode ini.")
    _h(d, "b. Laporan Pelaksanaan Pengasuransian BMN", 2)
    if polis_rows:
        _tabel(d, ["No", "Nomor Polis", "Penanggung", "Objek",
                   "Nilai Pertanggungan", "Periode"],
               [[str(i + 1), p_.get("nomor_polis") or "",
                 p_.get("penanggung") or "", p_.get("asset_name") or "",
                 fmt_rp(p_.get("nilai_pertanggungan")),
                 f"{p_.get('mulai') or ''} s.d. {p_.get('berakhir') or ''}"]
                for i, p_ in enumerate(polis_rows)], align_kanan={4})
    else:
        _p(d, "Nihil — tidak terdapat polis asuransi BMN tercatat.")
    _h(d, "c. Laporan BMN Berupa Rumah Negara", 2)
    if rumah_negara:
        from pembukuan_utils import parse_harga as _ph3
        _tabel(d, ["No", "Kode Barang", "Nama Barang", "NUP", "Nilai (Rp)"],
               [[str(i + 1), r_.get("asset_code") or "",
                 r_.get("asset_name") or "", r_.get("NUP") or "",
                 fmt_rp(_ph3(r_.get("purchase_price")))]
                for i, r_ in enumerate(rumah_negara)], align_kanan={4})
    else:
        _p(d, "Nihil — tidak terdapat BMN berupa rumah negara.")
    _h(d, "d. Berita Acara Rekonsiliasi Internal", 2)
    _p(d, "(Sisipkan BAR internal UAKPB–UAKPA periode ini.)")
    _h(d, "e. Neraca Percobaan dan Laporan Neraca", 2)
    _p(d, "(Sisipkan cetakan neraca percobaan berbasis akrual dan laporan "
          "neraca dari aplikasi SAKTI/sistem keuangan.)")
    _h(d, "f. Data Transfer Masuk dan Transfer Keluar", 2)
    if transfer_rows:
        _tabel(d, ["Kode", "Uraian", "Kode Barang", "Kuantitas", "Nilai (Rp)"],
               [[j.get("kode_transaksi") or "",
                 "Transfer Masuk" if str(j.get("kode_transaksi")) == "102"
                 else "Transfer Keluar",
                 j.get("kode_barang") or "", fmt_rp(j.get("jumlah")),
                 fmt_rp(j.get("nilai"))] for j in transfer_rows[:100]],
               align_kanan={3, 4})
    else:
        _p(d, "Nihil — tidak terdapat transfer masuk/keluar pada periode "
              "ini.")
    _h(d, "g. Rekapitulasi Transaksi Hibah", 2)
    if hibah_rows:
        _tabel(d, ["Kode", "Kode Barang", "Kuantitas", "Nilai (Rp)"],
               [[j.get("kode_transaksi") or "", j.get("kode_barang") or "",
                 fmt_rp(j.get("jumlah")), fmt_rp(j.get("nilai"))]
                for j in hibah_rows[:100]], align_kanan={2, 3})
    else:
        _p(d, "Nihil — tidak terdapat transaksi hibah pada periode ini.")
    _h(d, "h. Daftar BMN Hilang yang Diusulkan ke Pengelola", 2)
    if hilang_rows:
        _tabel(d, ["No", "Kode Barang", "NUP", "Nama Barang", "Status",
                   "Nomor SK"],
               [[str(i + 1), h_.get("asset_code") or "", h_.get("NUP") or "",
                 h_.get("asset_name") or "", h_.get("status") or "",
                 h_.get("nomor_sk") or ""] for i, h_ in
                enumerate(hilang_rows)])
    else:
        _p(d, "Nihil — tidak terdapat BMN hilang yang diusulkan "
              "penghapusannya.")
    _h(d, "i. Laporan Pengawasan dan Pengendalian", 2)
    _p(d, f"Register penertiban Wasdal berjalan: {n_tertib}. Rincian "
          "pengawasan & pengendalian tersedia pada Laporan Wasdal PMK "
          "207/2021 (dapat diunduh dari modul Wasdal aplikasi).")
    d.add_page_break()

    # PENUTUP + ttd
    _h(d, "PENUTUP", 1)
    _p(d, f"Demikian Laporan Barang Pengguna {nama_satker} {label} disusun "
          "sebagai wujud transparansi dan akuntabilitas pengelolaan BMN, "
          "untuk digunakan sebagai bahan penyusunan neraca serta bahan "
          "pengambilan keputusan pengelolaan BMN.")
    d.add_paragraph()
    _p(d, f"Posisi per {_fmt_tanggal_id(per_iso)}")
    _p(d, "Kuasa Pengguna Barang,", space_after=40)
    _p(d, ttd.get("nama") or "………………………………", bold=True, space_after=0)
    if ttd.get("nip"):
        _p(d, f"NIP. {ttd['nip']}", space_after=0)

    buf = io.BytesIO()
    d.save(buf)
    nama_file = (f"LBP_{(kode or 'satker')}_{tahun}"
                 f"{'_S' + str(sem) if sem else ''}.docx")
    return Response(
        content=buf.getvalue(),
        media_type=("application/vnd.openxmlformats-officedocument"
                    ".wordprocessingml.document"),
        headers={"Content-Disposition":
                 f'attachment; filename="{nama_file}"'})
