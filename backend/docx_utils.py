"""Helper bersama generator laporan **.docx (Word)** — versi *editable* dari
laporan resmi PDF.

Mencerminkan sistem desain ReportLab reports.py (kop surat, blok judul, tabel
identitas 2 kolom, tabel data bergaris, blok tanda tangan tim + Kuasa Pengguna
Barang, tembusan, nomor halaman) sehingga tiap laporan cukup menyusun
kontennya. Tujuannya: satker dapat MENYESUAIKAN naskah (narasi lokal, tindak
lanjut, dsb.) sebelum ditandatangani — sesuatu yang tidak bisa dilakukan pada
PDF final.

Aturan privasi TTD dari pegawai_utils dipatuhi: penandatangan Non-ASN — atau
nomor berformat NIK — TIDAK dicetak NIP/NIK-nya di area tanda tangan
(`baris_identitas_ttd`), sama seperti laporan PDF.

Fungsi murni terhadap Mongo: peta status kepegawaian per-NIP disiapkan pemanggil
(mis. reports._peta_status_kepegawaian) lalu diteruskan ke signature_block.
"""
import base64
import io

_NAVY = "1F4E79"          # biru tua kop/heading (samakan dgn LBP)
_HEADER_BG = "1F4E79"     # latar header tabel
_ZEBRA = "F2F6FB"         # baris selang-seling


# ---------------------------------------------------------------------------
# Dasar dokumen & elemen kecil
# ---------------------------------------------------------------------------

def doc_baru(margin_cm=2.2, font="Cambria", size=11):
    """Document A4 baru dengan margin & font standar laporan."""
    from docx import Document
    from docx.shared import Cm, Pt
    from docx.enum.section import WD_ORIENT
    d = Document()
    st = d.styles["Normal"]
    st.font.name = font
    st.font.size = Pt(size)
    for s in d.sections:
        s.orientation = WD_ORIENT.PORTRAIT
        s.page_height = Cm(29.7)
        s.page_width = Cm(21.0)
        s.top_margin = Cm(margin_cm)
        s.bottom_margin = Cm(margin_cm)
        s.left_margin = Cm(margin_cm)
        s.right_margin = Cm(margin_cm)
    return d


def _set_shading(cell, hexcolor):
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:fill"), hexcolor)
    cell._tc.get_or_add_tcPr().append(shd)


def _no_cell_margins_tight(cell):
    """Padding sel tipis agar tabel rapat (opsional)."""
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    tcPr = cell._tc.get_or_add_tcPr()
    m = OxmlElement("w:tcMar")
    for tag, val in (("top", 20), ("bottom", 20), ("start", 60), ("end", 60)):
        e = OxmlElement(f"w:{tag}")
        e.set(qn("w:w"), str(val))
        e.set(qn("w:type"), "dxa")
        m.append(e)
    tcPr.append(m)


def page_footer(d, teks=""):
    """Footer: nama laporan (kiri, kecil abu) + 'Halaman X' (kanan, field PAGE)."""
    from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT
    from docx.shared import Pt, RGBColor, Cm
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    for s in d.sections:
        p = s.footer.paragraphs[0]
        p.text = ""
        # tab kanan di batas margin kanan
        usable = s.page_width - s.left_margin - s.right_margin
        p.paragraph_format.tab_stops.add_tab_stop(usable, WD_TAB_ALIGNMENT.RIGHT)
        r = p.add_run((teks or "") + "\t")
        r.font.size = Pt(7.5)
        r.font.color.rgb = RGBColor(0x80, 0x80, 0x80)
        r2 = p.add_run("Halaman ")
        r2.font.size = Pt(7.5)
        r2.font.color.rgb = RGBColor(0x80, 0x80, 0x80)
        fld = OxmlElement("w:fldSimple")
        fld.set(qn("w:instr"), "PAGE")
        p._p.append(fld)


def para(d, teks, *, bold=False, italic=False, center=False, justify=True,
         size=None, space_after=6, space_before=0):
    """Paragraf teks biasa (justify default seperti badan laporan)."""
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Pt
    p = d.add_paragraph()
    run = p.add_run(teks or "")
    run.bold = bold
    run.italic = italic
    if size:
        run.font.size = Pt(size)
    if center:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    elif justify:
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.space_before = Pt(space_before)
    return p


def _garis_bawah(p, sz=6, color=_NAVY):
    """Border bawah pada paragraf (untuk garis kop)."""
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    pPr = p._p.get_or_add_pPr()
    pbdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), str(sz))
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), color)
    pbdr.append(bottom)
    pPr.append(pbdr)


# ---------------------------------------------------------------------------
# Kop surat
# ---------------------------------------------------------------------------

def kop_surat(d, settings):
    """Kop surat: logo (bila data-URI) di atas-tengah, lalu nama instansi
    (besar), unit organisasi + sub-unit (tebal, kapital), alamat (kecil),
    ditutup garis ganda. Semua terpusat — mudah disunting di Word."""
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Cm, Pt, RGBColor
    settings = settings or {}
    nama_instansi = str(settings.get("nama_instansi") or "").strip()
    nama_unit = str(settings.get("nama_unit_organisasi") or "").strip()
    nama_sub = str(settings.get("nama_sub_unit") or "").strip()
    alamat = str(settings.get("alamat_instansi") or "").strip()
    logo_url = str(settings.get("logo_url") or "")

    if logo_url.startswith("data:"):
        try:
            _hdr, b64 = logo_url.split(",", 1)
            buf = io.BytesIO(base64.b64decode(b64))
            p = d.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.add_run().add_picture(buf, height=Cm(1.9))
            p.paragraph_format.space_after = Pt(2)
        except Exception:
            pass

    def _baris(teks, sz, bold, upper=False, color=None, after=1):
        if not teks:
            return
        p = d.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(teks.upper() if upper else teks)
        r.bold = bold
        r.font.size = Pt(sz)
        if color:
            r.font.color.rgb = color
        p.paragraph_format.space_after = Pt(after)
        p.paragraph_format.space_before = Pt(0)
        return p

    _baris(nama_instansi, 13, False, upper=True)
    _baris(nama_unit, 12, True, upper=True)
    _baris(nama_sub, 12, True, upper=True)
    last = None
    for ln in alamat.splitlines():
        if ln.strip():
            last = _baris(ln.strip(), 9, False, color=RGBColor(0x33, 0x33, 0x33))
    if last is None:
        # tetap sediakan paragraf pembawa garis meski alamat kosong
        last = d.add_paragraph()
    _garis_bawah(last, sz=18)
    d.add_paragraph().paragraph_format.space_after = Pt(2)


def title_block(d, judul, nomor=None, subjudul=None):
    """Blok judul dokumen (tengah, tebal) + sub-judul + 'Nomor: ...'."""
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Pt, RGBColor
    for i, baris in enumerate((judul or "").split("\n")):
        p = d.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(baris.strip())
        r.bold = True
        r.font.size = Pt(13)
        r.font.color.rgb = RGBColor(0x1F, 0x2A, 0x44)
        p.paragraph_format.space_after = Pt(1)
        p.paragraph_format.space_before = Pt(0)
    if subjudul:
        para(d, subjudul, center=True, justify=False, bold=True, size=11, space_after=1)
    if nomor:
        para(d, f"Nomor: {nomor}", center=True, justify=False, size=10.5, space_after=6)
    else:
        d.add_paragraph().paragraph_format.space_after = Pt(4)


def section(d, judul, nomor_romawi=None):
    """Judul bagian tebal, mis. 'I. SUSUNAN TIM'."""
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    p = d.add_paragraph()
    teks = f"{nomor_romawi}. {judul}" if nomor_romawi else judul
    r = p.add_run(teks)
    r.bold = True
    r.font.size = Pt(11)
    r.font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(3)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    return p


# ---------------------------------------------------------------------------
# Tabel
# ---------------------------------------------------------------------------

def meta_table(d, rows, label_w_cm=4.2):
    """Tabel identitas 2 kolom tanpa garis: 'Label : Nilai' (blok 'Yang
    bertanda tangan di bawah ini')."""
    from docx.shared import Cm, Pt
    t = d.add_table(rows=0, cols=3)
    for label, val in rows:
        cells = t.add_row().cells
        cells[0].text = ""
        cells[0].paragraphs[0].add_run(str(label)).font.size = Pt(11)
        cells[1].text = ":"
        cells[2].text = ""
        r = cells[2].paragraphs[0].add_run(str(val or "-"))
        r.font.size = Pt(11)
    # lebar kolom
    for row in t.rows:
        row.cells[0].width = Cm(label_w_cm)
        row.cells[1].width = Cm(0.5)
    return t


def data_table(d, header, rows, *, widths_cm=None, align_right=None,
               align_center=None, zebra=True, font_size=9):
    """Tabel data bergaris: header biru-putih tebal; align_right/align_center =
    himpunan indeks kolom; zebra = selang-seling latar baris."""
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Cm, Pt, RGBColor
    ncol = len(header)
    t = d.add_table(rows=1, cols=ncol)
    t.style = "Table Grid"
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    kanan = set(align_right or ())
    tengah = set(align_center or ())
    for i, judul in enumerate(header):
        c = t.rows[0].cells[i]
        c.text = ""
        p = c.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(str(judul))
        r.bold = True
        r.font.size = Pt(font_size)
        r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        _set_shading(c, _HEADER_BG)
    for ridx, row in enumerate(rows):
        cells = t.add_row().cells
        for i, val in enumerate(row):
            c = cells[i]
            c.text = ""
            p = c.paragraphs[0]
            if i in kanan:
                p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            elif i in tengah:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            else:
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            r = p.add_run("" if val is None else str(val))
            r.font.size = Pt(font_size)
            if zebra and ridx % 2 == 1:
                _set_shading(c, _ZEBRA)
    if widths_cm and len(widths_cm) == ncol:
        for row in t.rows:
            for i, w in enumerate(widths_cm):
                row.cells[i].width = Cm(w)
    return t


# ---------------------------------------------------------------------------
# Blok tanda tangan (tim + Kuasa Pengguna Barang) — kaidah BA inventarisasi
# ---------------------------------------------------------------------------

def _sig_cell(cell, header, role, nama, nomor_baris, pre=None):
    """Isi satu sel tanda tangan: (pre) + header + peran + ruang ttd + nama +
    baris NIP. `pre` = baris di ATAS header (mis. 'Dibuat di …', 'tanggal …')."""
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Pt
    cell.text = ""
    def _ln(teks, bold=False, underline=False, size=10.5, after=0):
        p = cell.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(teks or "")
        r.bold = bold
        r.underline = underline
        r.font.size = Pt(size)
        p.paragraph_format.space_after = Pt(after)
        p.paragraph_format.space_before = Pt(0)
        return p
    # buang paragraf kosong bawaan sel
    cell.paragraphs[0]._p.getparent().remove(cell.paragraphs[0]._p)
    for baris in (pre or []):
        _ln(baris)
    if header:
        _ln(header)
    if role:
        _ln(role)
    _ln("")  # ruang tanda tangan
    _ln("")
    _ln("")
    _ln(nama or "________________________", bold=True, underline=bool(nama))
    if nomor_baris:
        _ln(nomor_baris, size=10)


def signature_block(d, tim, ident, tempat_tanggal, *, label_tim="Tim",
                    status_by_nip=None, header_mengetahui="Mengetahui,",
                    saksi=None):
    """Blok TTD: SELURUH anggota `tim` bertanda tangan (ketua ditandai;
    berpasangan 2 kolom per baris; tempat/tanggal di kanan atas), lalu
    (opsional) saksi, lalu Kuasa Pengguna Barang 'Mengetahui' di TENGAH bawah.
    Non-ASN/NIK: baris NIP/NIK tidak dicetak."""
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Emu, Pt
    from pegawai_utils import baris_identitas_laporan, baris_identitas_ttd
    peta = status_by_nip or {}

    def _member(m):
        return m if isinstance(m, dict) else {"nama": str(m).strip() or "-"}

    anggota = [_member(m) for m in (tim or [])]
    if not anggota:
        anggota = [{"nama": "", "nip": ""}]
    punya_ketua = any(m.get("is_ketua") for m in anggota)

    p = d.add_paragraph()
    r = p.add_run(f"{label_tim}:")
    r.bold = True
    r.font.size = Pt(11)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(2)

    usable = int(d.sections[0].page_width - d.sections[0].left_margin - d.sections[0].right_margin)
    half = Emu(usable // 2)

    for i in range(0, len(anggota), 2):
        pasangan = anggota[i:i + 2]
        t = d.add_table(rows=1, cols=2)
        t.alignment = WD_ALIGN_PARAGRAPH.CENTER
        # tempat/tanggal di kanan atas hanya pada baris pertama
        if i == 0:
            head = t.rows[0].cells
            head[0].text = ""
            hp = head[1].paragraphs[0]
            hp.alignment = WD_ALIGN_PARAGRAPH.CENTER
            hp.add_run(tempat_tanggal or "").font.size = Pt(10.5)
        else:
            t.rows[0].cells[0].text = ""
            t.rows[0].cells[1].text = ""
        cells = t.add_row().cells
        for j in range(2):
            if j < len(pasangan):
                m = pasangan[j]
                idx = i + j
                ketua = m.get("is_ketua") if punya_ketua else (idx == 0)
                nip_m = str(m.get("nip") or "").strip()
                baris = baris_identitas_ttd(nip_m, "NIP. ........................",
                                            peta.get(nip_m, ""))
                _sig_cell(cells[j], "Ketua Tim," if ketua else "Anggota,", None,
                          m.get("nama"), baris[0] if baris else "")
            else:
                cells[j].text = ""
        for row in t.rows:
            row.cells[0].width = half
            row.cells[1].width = half
        d.add_paragraph().paragraph_format.space_after = Pt(2)

    # Saksi (opsional)
    for s in (saksi or []):
        sm = _member(s)
        nip_s = str(sm.get("nip") or "").strip()
        baris = baris_identitas_ttd(nip_s, "", peta.get(nip_s, ""))
        t = d.add_table(rows=1, cols=2)
        cells = t.rows[0].cells
        cells[1].text = ""
        _sig_cell(cells[0], "Saksi,", sm.get("jabatan") or None, sm.get("nama"),
                  baris[0] if baris else "")
        for row in t.rows:
            row.cells[0].width = half
            row.cells[1].width = half

    # Kuasa Pengguna Barang — Mengetahui, tengah bawah
    nip_kpb = str(ident.get("kasatker_nip") or "").strip()
    baris_kpb = (baris_identitas_laporan(nip_kpb, peta.get(nip_kpb, ""))
                 if nip_kpb else "NIP. ")
    t = d.add_table(rows=1, cols=3)
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cells = t.rows[0].cells
    cells[0].text = ""
    cells[2].text = ""
    _sig_cell(cells[1], header_mengetahui,
              str(ident.get("kasatker_jabatan") or "Kuasa Pengguna Barang,"),
              ident.get("kasatker_nama"), baris_kpb)
    third = Emu(usable // 3)
    for i in (0, 1, 2):
        t.rows[0].cells[i].width = third


def signature_single(d, *, nama, header="Yang membuat pernyataan,", jabatan=None,
                     pre_lines=None, nip="", status=""):
    """Blok tanda tangan TUNGGAL (mis. SPTJM/Surat Koreksi: Kuasa Pengguna
    Barang) di sisi KANAN. Baris NIP mengikuti aturan Non-ASN/NIK
    (baris_identitas_laporan). `pre_lines` = tempat/tanggal di atas header."""
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Emu
    from pegawai_utils import baris_identitas_laporan
    nomor_baris = baris_identitas_laporan(str(nip or "").strip(), status) if str(nip or "").strip() else "NIP. "
    usable = int(d.sections[0].page_width - d.sections[0].left_margin - d.sections[0].right_margin)
    t = d.add_table(rows=1, cols=2)
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cells = t.rows[0].cells
    cells[0].text = ""
    _sig_cell(cells[1], header, jabatan, nama, nomor_baris, pre=pre_lines)
    cells[0].width = Emu(usable // 2)
    cells[1].width = Emu(usable // 2)


def identity_block(d, rows, *, intro="Yang bertanda tangan di bawah ini:"):
    """Intro + tabel identitas 'Label : Nilai' (blok pembuka surat pernyataan)."""
    if intro:
        para(d, intro, justify=False, space_after=2)
    meta_table(d, rows)


def tembusan(d, lines):
    """Blok 'Tembusan:' berpoin (bila ada)."""
    from docx.shared import Pt
    baris = [ln.strip() for ln in (lines or []) if str(ln).strip()]
    if not baris:
        return
    p = d.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    r = p.add_run("Tembusan:")
    r.bold = True
    r.font.size = Pt(9)
    for i, ln in enumerate(baris, 1):
        pp = d.add_paragraph()
        pp.paragraph_format.space_after = Pt(0)
        rr = pp.add_run(f"{i}. {ln}")
        rr.font.size = Pt(9)


def to_bytes(d):
    """Serialisasi Document → bytes .docx."""
    buf = io.BytesIO()
    d.save(buf)
    buf.seek(0)
    return buf.getvalue()
