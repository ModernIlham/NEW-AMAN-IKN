"""Jaring pengaman `docx_utils` — 17 fungsi yang memproduksi dokumen Word resmi.

Modul ini adalah sistem desain bersama untuk seluruh laporan .docx (BA
inventarisasi, SPTJM, surat koreksi, DBHI). Sampai kini ia berjalan TANPA satu
pun uji: sebuah kesalahan kecil di sini menghasilkan berita acara yang salah
cetak di seluruh satker, dan baru ketahuan setelah dokumen ditandatangani.

Yang paling dijaga di berkas ini:

  • **Privasi area tanda tangan.** NIK dan penandatangan Non-ASN TIDAK boleh
    tercetak nomornya (aturan pemilik, diwarisi dari laporan PDF). Regresi di
    jalur ini membocorkan NIK ke dokumen resmi yang beredar.
  • **Semua anggota tim dapat kolom tanda tangan** — bukan hanya dua yang
    pertama. Berkas ini menyusunnya berpasangan; kesalahan iterasi akan
    "menghilangkan" anggota tanpa error apa pun.
  • **Dokumen tetap terbentuk** meski masukan cacat (logo rusak, alamat
    kosong, settings None) — generator laporan tidak boleh 500 karena kop.

Uji membaca ULANG hasil `to_bytes()` lewat python-docx, jadi dokumen yang
korup akan gagal dibuka, bukan lolos diam-diam.
"""
import io
import zipfile

import pytest
from docx import Document
from docx.oxml.ns import qn
from docx.shared import Cm, Pt

import docx_utils as dx


# ---------------------------------------------------------------------------
# Bantuan baca-ulang
# ---------------------------------------------------------------------------

def _buka_ulang(d):
    """Serialisasi lalu buka lagi — memastikan XML-nya benar-benar sah."""
    return Document(io.BytesIO(dx.to_bytes(d)))


def _xml_dokumen(d):
    """Isi `word/document.xml` sebagai teks.

    .docx adalah ZIP TERKOMPRESI: mencari string pada `to_bytes()` mentah
    selalu "tidak ketemu" sehingga uji privasi lolos tanpa membuktikan apa
    pun. Buka arsipnya dulu.
    """
    with zipfile.ZipFile(io.BytesIO(dx.to_bytes(d))) as z:
        return z.read("word/document.xml").decode("utf-8")


def _dekat(a, b, toleransi=None):
    """Ukuran Word disimpan dalam twips; pembulatan EMU→twips→EMU membuat
    perbandingan persis meleset beberapa ratus EMU (< 0,001 cm)."""
    return abs(int(a) - int(b)) <= int(toleransi if toleransi is not None else Cm(0.01))


def _run_berisi(p):
    """Run pertama yang benar-benar membawa teks — `cell.text = ""` menyisakan
    satu run kosong di depan."""
    return next((r for r in p.runs if r.text), None)


def _baris_sel(cell):
    """Teks per paragraf dalam satu sel (baris kosong ikut, karena ruang tanda
    tangan memang berupa paragraf kosong)."""
    return [p.text for p in cell.paragraphs]


def _isi_sel(cell):
    """Hanya baris berisi — untuk menguji urutan label tanpa terganggu ruang
    tanda tangan."""
    return [t for t in _baris_sel(cell) if t.strip()]


def _fill(cell):
    """Warna latar sel (`w:shd/@w:fill`) atau None bila tak diwarnai."""
    tcPr = cell._tc.tcPr
    if tcPr is None:
        return None
    shd = tcPr.find(qn("w:shd"))
    return None if shd is None else shd.get(qn("w:fill"))


def _punya_garis_bawah(p):
    return p._p.find(qn("w:pPr") + "/" + qn("w:pBdr")) is not None or bool(
        p._p.xpath(".//w:pBdr/w:bottom"))


def _teks_paragraf(d):
    return [p.text for p in d.paragraphs]


# ---------------------------------------------------------------------------
# Dasar dokumen
# ---------------------------------------------------------------------------

class TestDokumenDasar:
    def test_potret_a4_dan_margin(self):
        d = dx.doc_baru(margin_cm=2.2)
        s = d.sections[0]
        assert _dekat(s.page_width, Cm(21.0))
        assert _dekat(s.page_height, Cm(29.7))
        assert _dekat(s.top_margin, Cm(2.2)) and _dekat(s.bottom_margin, Cm(2.2))
        assert _dekat(s.left_margin, Cm(2.2)) and _dekat(s.right_margin, Cm(2.2))

    def test_landscape_menukar_sisi(self):
        """DBHI/DBKP bergantung pada ini. Bila penukaran hilang, tabel lebar
        tercetak terpotong di kertas potret — tanpa error apa pun."""
        from docx.enum.section import WD_ORIENT
        d = dx.doc_baru(landscape=True)
        s = d.sections[0]
        assert _dekat(s.page_width, Cm(29.7))
        assert _dekat(s.page_height, Cm(21.0))
        assert s.page_width > s.page_height
        assert s.orientation == WD_ORIENT.LANDSCAPE

    def test_font_dan_ukuran_normal(self):
        d = dx.doc_baru(font="Cambria", size=11)
        st = d.styles["Normal"]
        assert st.font.name == "Cambria"
        assert st.font.size == Pt(11)

    def test_margin_bisa_dikecilkan(self):
        d = dx.doc_baru(margin_cm=1.0)
        assert _dekat(d.sections[0].left_margin, Cm(1.0))

    def test_to_bytes_menghasilkan_docx_yang_bisa_dibuka(self):
        d = dx.doc_baru()
        dx.para(d, "isi")
        b = dx.to_bytes(d)
        assert b[:2] == b"PK", "docx = zip; awalan bukan PK berarti berkas rusak"
        assert len(b) > 1000
        assert _teks_paragraf(Document(io.BytesIO(b))) == ["isi"]

    def test_to_bytes_tidak_menyisakan_posisi_baca(self):
        """`seek(0)` sebelum `getvalue()` mudah terhapus saat refactor; tanpa
        itu berkas terkirim kosong ke peramban."""
        d = dx.doc_baru()
        assert len(dx.to_bytes(d)) == len(dx.to_bytes(d))


# ---------------------------------------------------------------------------
# Footer
# ---------------------------------------------------------------------------

class TestFooter:
    def test_nomor_halaman_memakai_field_page(self):
        """Angka halaman WAJIB berupa field PAGE, bukan teks mati — dokumen
        ini disunting satker, jumlah halamannya berubah."""
        d = dx.doc_baru()
        dx.page_footer(d, "Berita Acara")
        d2 = _buka_ulang(d)
        p = d2.sections[0].footer.paragraphs[0]
        xml = p._p.xml
        assert "fldSimple" in xml
        assert 'w:instr="PAGE"' in xml
        assert "Halaman" in p.text

    def test_teks_footer_muncul(self):
        d = dx.doc_baru()
        dx.page_footer(d, "Laporan Hasil Inventarisasi")
        assert "Laporan Hasil Inventarisasi" in _buka_ulang(d).sections[0].footer.paragraphs[0].text

    def test_teks_kosong_tetap_menomori_halaman(self):
        d = dx.doc_baru()
        dx.page_footer(d)
        p = _buka_ulang(d).sections[0].footer.paragraphs[0]
        assert "Halaman" in p.text
        assert 'w:instr="PAGE"' in p._p.xml

    def test_tab_kanan_di_batas_margin(self):
        d = dx.doc_baru(margin_cm=2.0)
        dx.page_footer(d, "x")
        s = d.sections[0]
        stops = s.footer.paragraphs[0].paragraph_format.tab_stops
        assert len(stops) == 1
        assert stops[0].position == s.page_width - s.left_margin - s.right_margin


# ---------------------------------------------------------------------------
# Paragraf
# ---------------------------------------------------------------------------

class TestParagraf:
    def test_teks_none_tidak_meledak(self):
        d = dx.doc_baru()
        dx.para(d, None)
        assert _teks_paragraf(d) == [""]

    def test_rata_kanan_kiri_bawaan(self):
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        d = dx.doc_baru()
        p = dx.para(d, "badan laporan")
        assert p.alignment == WD_ALIGN_PARAGRAPH.JUSTIFY

    def test_center_mengalahkan_justify(self):
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        d = dx.doc_baru()
        p = dx.para(d, "judul", center=True)
        assert p.alignment == WD_ALIGN_PARAGRAPH.CENTER

    def test_justify_false_membiarkan_rata_kiri(self):
        d = dx.doc_baru()
        assert dx.para(d, "x", justify=False).alignment is None

    def test_gaya_dan_ukuran_terpasang(self):
        d = dx.doc_baru()
        r = dx.para(d, "x", bold=True, italic=True, size=9).runs[0]
        assert r.bold is True and r.italic is True
        assert r.font.size == Pt(9)

    def test_jarak_atas_bawah(self):
        d = dx.doc_baru()
        p = dx.para(d, "x", space_after=12, space_before=4)
        assert p.paragraph_format.space_after == Pt(12)
        assert p.paragraph_format.space_before == Pt(4)


# ---------------------------------------------------------------------------
# Kop surat
# ---------------------------------------------------------------------------

_KOP = {
    "nama_instansi": "Otorita Ibu Kota Nusantara",
    "nama_unit_organisasi": "Deputi Bidang Sarana dan Prasarana",
    "nama_sub_unit": "Direktorat Barang Milik Negara",
    "alamat_instansi": "Jalan Sumbu Kebangsaan No. 1\nNusantara 76111",
}


class TestKopSurat:
    def test_instansi_dan_unit_dikapitalkan(self):
        d = dx.doc_baru()
        dx.kop_surat(d, _KOP)
        teks = _teks_paragraf(d)
        assert "OTORITA IBU KOTA NUSANTARA" in teks
        assert "DEPUTI BIDANG SARANA DAN PRASARANA" in teks
        assert "DIREKTORAT BARANG MILIK NEGARA" in teks

    def test_alamat_tidak_dikapitalkan(self):
        d = dx.doc_baru()
        dx.kop_surat(d, _KOP)
        assert "Jalan Sumbu Kebangsaan No. 1" in _teks_paragraf(d)

    def test_alamat_multibaris_jadi_paragraf_terpisah(self):
        d = dx.doc_baru()
        dx.kop_surat(d, _KOP)
        teks = _teks_paragraf(d)
        assert teks.index("Jalan Sumbu Kebangsaan No. 1") + 1 == teks.index("Nusantara 76111")

    def test_baris_alamat_kosong_dilewati(self):
        d = dx.doc_baru()
        dx.kop_surat(d, {"nama_instansi": "X", "alamat_instansi": "A\n\n  \nB"})
        assert [t for t in _teks_paragraf(d) if t] == ["X", "A", "B"]

    def test_alamat_kosong_tetap_bergaris(self):
        """Tanpa paragraf pembawa cadangan, kop tanpa alamat kehilangan garis
        pemisahnya — atau lebih buruk, melempar AttributeError."""
        d = dx.doc_baru()
        dx.kop_surat(d, {"nama_instansi": "Otorita"})
        assert any(_punya_garis_bawah(p) for p in d.paragraphs)

    def test_garis_bawah_terpasang_saat_alamat_ada(self):
        d = dx.doc_baru()
        dx.kop_surat(d, _KOP)
        bergaris = [p.text for p in d.paragraphs if _punya_garis_bawah(p)]
        assert bergaris == ["Nusantara 76111"], "garis harus di baris alamat TERAKHIR"

    def test_logo_rusak_tidak_menggagalkan_dokumen(self):
        """Logo satker tersimpan sebagai data-URI di setelan; satu karakter
        rusak tak boleh membuat seluruh laporan gagal diunduh."""
        d = dx.doc_baru()
        dx.kop_surat(d, dict(_KOP, logo_url="data:image/png;base64,INI-BUKAN-BASE64!!"))
        assert "OTORITA IBU KOTA NUSANTARA" in _teks_paragraf(d)
        assert dx.to_bytes(d)[:2] == b"PK"

    def test_logo_bukan_data_uri_diabaikan(self):
        d = dx.doc_baru()
        dx.kop_surat(d, dict(_KOP, logo_url="https://contoh.id/logo.png"))
        assert dx.to_bytes(d)[:2] == b"PK"

    def test_settings_none_aman(self):
        d = dx.doc_baru()
        dx.kop_surat(d, None)
        assert dx.to_bytes(d)[:2] == b"PK"

    def test_semua_baris_kop_terpusat(self):
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        d = dx.doc_baru()
        dx.kop_surat(d, _KOP)
        berisi = [p for p in d.paragraphs if p.text]
        assert berisi and all(p.alignment == WD_ALIGN_PARAGRAPH.CENTER for p in berisi)


# ---------------------------------------------------------------------------
# Judul & bagian
# ---------------------------------------------------------------------------

class TestJudulDanBagian:
    def test_judul_multibaris_jadi_paragraf_sendiri(self):
        d = dx.doc_baru()
        dx.title_block(d, "BERITA ACARA\nHASIL INVENTARISASI")
        assert _teks_paragraf(d)[:2] == ["BERITA ACARA", "HASIL INVENTARISASI"]

    def test_judul_tebal_dan_terpusat(self):
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        d = dx.doc_baru()
        dx.title_block(d, "BERITA ACARA")
        p = d.paragraphs[0]
        assert p.alignment == WD_ALIGN_PARAGRAPH.CENTER
        assert p.runs[0].bold is True

    def test_nomor_dan_subjudul(self):
        d = dx.doc_baru()
        dx.title_block(d, "BERITA ACARA", nomor="BA-12/OIKN/2026", subjudul="Barang Tidak Ditemukan")
        teks = _teks_paragraf(d)
        assert "Barang Tidak Ditemukan" in teks
        assert "Nomor: BA-12/OIKN/2026" in teks
        assert teks.index("Barang Tidak Ditemukan") < teks.index("Nomor: BA-12/OIKN/2026")

    def test_tanpa_nomor_tidak_mencetak_label_kosong(self):
        d = dx.doc_baru()
        dx.title_block(d, "BERITA ACARA")
        assert not any(t.startswith("Nomor:") for t in _teks_paragraf(d))

    def test_judul_none_tidak_meledak(self):
        d = dx.doc_baru()
        dx.title_block(d, None)
        assert dx.to_bytes(d)[:2] == b"PK"

    def test_section_dengan_romawi(self):
        d = dx.doc_baru()
        dx.section(d, "SUSUNAN TIM", "I")
        assert d.paragraphs[0].text == "I. SUSUNAN TIM"
        assert d.paragraphs[0].runs[0].bold is True

    def test_section_tanpa_romawi(self):
        d = dx.doc_baru()
        dx.section(d, "PENUTUP")
        assert d.paragraphs[0].text == "PENUTUP"


# ---------------------------------------------------------------------------
# Tabel
# ---------------------------------------------------------------------------

class TestMetaTable:
    def test_bentuk_tiga_kolom_dengan_titik_dua(self):
        d = dx.doc_baru()
        t = dx.meta_table(d, [("Nama", "Budi"), ("Jabatan", "Ketua Tim")])
        assert len(t.rows) == 2 and len(t.columns) == 3
        assert [c.text for c in t.rows[0].cells] == ["Nama", ":", "Budi"]

    def test_nilai_kosong_jadi_strip(self):
        """Sel benar-benar kosong membuat tabel identitas tampak salah cetak."""
        d = dx.doc_baru()
        t = dx.meta_table(d, [("NIP", None), ("NIK", ""), ("Unit", 0)])
        assert [t.rows[i].cells[2].text for i in range(3)] == ["-", "-", "-"]

    def test_nilai_angka_bukan_nol_tetap_tercetak(self):
        d = dx.doc_baru()
        t = dx.meta_table(d, [("Jumlah", 12)])
        assert t.rows[0].cells[2].text == "12"

    def test_lebar_kolom_label(self):
        d = dx.doc_baru()
        t = dx.meta_table(d, [("A", "B")], label_w_cm=5.0)
        assert _dekat(t.rows[0].cells[0].width, Cm(5.0))
        assert _dekat(t.rows[0].cells[1].width, Cm(0.5))

    def test_baris_kosong_menghasilkan_tabel_kosong(self):
        d = dx.doc_baru()
        assert len(dx.meta_table(d, []).rows) == 0


class TestDataTable:
    _H = ["No", "Kode Barang", "Nilai"]
    _R = [[1, "3.05.01.01.001", 1500000], [2, "3.05.01.01.002", None],
          [3, "3.05.01.01.003", 250000]]

    def test_header_plus_semua_baris(self):
        d = dx.doc_baru()
        t = dx.data_table(d, self._H, self._R)
        assert len(t.rows) == 1 + len(self._R)
        assert [c.text for c in t.rows[0].cells] == self._H

    def test_header_putih_tebal_berlatar_navy(self):
        from docx.shared import RGBColor
        d = dx.doc_baru()
        t = dx.data_table(d, self._H, self._R)
        for c in t.rows[0].cells:
            assert _fill(c) == "1F4E79"
            r = _run_berisi(c.paragraphs[0])
            assert r is not None and r.bold is True
            assert r.font.color.rgb == RGBColor(0xFF, 0xFF, 0xFF), \
                "teks header di atas latar navy WAJIB putih agar terbaca"

    def test_zebra_hanya_baris_ganjil(self):
        """Selang-seling salah fase = seluruh tabel tampak diarsir."""
        d = dx.doc_baru()
        t = dx.data_table(d, self._H, self._R, zebra=True)
        assert _fill(t.rows[1].cells[0]) is None
        assert _fill(t.rows[2].cells[0]) == "F2F6FB"
        assert _fill(t.rows[3].cells[0]) is None

    def test_zebra_bisa_dimatikan(self):
        d = dx.doc_baru()
        t = dx.data_table(d, self._H, self._R, zebra=False)
        assert all(_fill(t.rows[i].cells[0]) is None for i in (1, 2, 3))

    def test_sel_none_jadi_kosong_bukan_teks_none(self):
        d = dx.doc_baru()
        t = dx.data_table(d, self._H, self._R)
        assert t.rows[2].cells[2].text == ""

    def test_perataan_kolom(self):
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        d = dx.doc_baru()
        t = dx.data_table(d, self._H, self._R, align_right={2}, align_center={0})
        sel = t.rows[1].cells
        assert sel[0].paragraphs[0].alignment == WD_ALIGN_PARAGRAPH.CENTER
        assert sel[1].paragraphs[0].alignment == WD_ALIGN_PARAGRAPH.LEFT
        assert sel[2].paragraphs[0].alignment == WD_ALIGN_PARAGRAPH.RIGHT

    def test_lebar_kolom_dipasang_bila_jumlahnya_cocok(self):
        d = dx.doc_baru()
        t = dx.data_table(d, self._H, self._R, widths_cm=[1.0, 6.0, 3.0])
        for baris in t.rows:   # header ikut, kalau tidak kolomnya tak sejajar
            for c, w in zip(baris.cells, (1.0, 6.0, 3.0)):
                assert _dekat(c.width, Cm(w))

    def test_lebar_kolom_diabaikan_bila_jumlahnya_beda(self):
        """Menerapkan daftar lebar yang lebih pendek akan melempar IndexError
        di tengah pembuatan laporan; membiarkan lebar bawaan jauh lebih aman."""
        d = dx.doc_baru()
        t = dx.data_table(d, self._H, self._R, widths_cm=[1.0, 6.0])
        lebar = [c.width for c in t.rows[1].cells]
        assert len(set(lebar)) == 1, "kolom harus tetap di lebar bawaan yang seragam"
        assert not _dekat(lebar[0], Cm(1.0))

    def test_tanpa_baris_data_tetap_bergaris(self):
        d = dx.doc_baru()
        t = dx.data_table(d, self._H, [])
        assert len(t.rows) == 1
        assert t.style.name == "Table Grid"

    def test_baris_lebih_pendek_dari_header_ditoleransi(self):
        d = dx.doc_baru()
        t = dx.data_table(d, self._H, [[1, "kode"]])
        assert [c.text for c in t.rows[1].cells] == ["1", "kode", ""]


# ---------------------------------------------------------------------------
# Blok tanda tangan — jalur paling sensitif
# ---------------------------------------------------------------------------

_NIP = "198001012005011001"
_NIP2 = "199002022010012002"
_NIK = "3506042503900001"
_IDENT = {"kasatker_nama": "Dr. Sartono", "kasatker_nip": "197001011990031002",
          "kasatker_jabatan": "Kuasa Pengguna Barang,"}


def _sel_ttd(d):
    """Semua sel tanda tangan anggota (baris kedua tiap tabel berpasangan)."""
    keluar = []
    for t in d.tables:
        for r in t.rows:
            for c in r.cells:
                if c.text.strip():
                    keluar.append(_isi_sel(c))
    return keluar


class TestBlokTandaTangan:
    def test_semua_anggota_dapat_kolom(self):
        """Lima anggota → lima kolom. Iterasi yang salah "menghilangkan"
        anggota tanpa error; berita acara lalu kurang penanda tangan."""
        tim = [{"nama": f"Anggota {i}", "nip": ""} for i in range(5)]
        d = dx.doc_baru()
        dx.signature_block(d, tim, _IDENT, "Nusantara, 1 Januari 2026")
        semua = "\n".join(p.text for t in d.tables for r in t.rows for c in r.cells
                          for p in c.paragraphs)
        for i in range(5):
            assert f"Anggota {i}" in semua

    def test_dua_kolom_per_baris(self):
        tim = [{"nama": f"P{i}"} for i in range(4)]
        d = dx.doc_baru()
        dx.signature_block(d, tim, _IDENT, "Nusantara, 1 Januari 2026")
        # 4 anggota → 2 tabel pasangan + 1 tabel KPB
        assert len(d.tables) == 3
        assert all(len(t.columns) == 2 for t in d.tables[:2])

    def test_anggota_ganjil_menyisakan_sel_kosong(self):
        d = dx.doc_baru()
        dx.signature_block(d, [{"nama": "A"}, {"nama": "B"}, {"nama": "C"}],
                           _IDENT, "Nusantara, 1 Januari 2026")
        assert d.tables[1].rows[1].cells[1].text.strip() == ""

    def test_ketua_eksplisit_menang_atas_urutan(self):
        d = dx.doc_baru()
        dx.signature_block(d, [{"nama": "A"}, {"nama": "B", "is_ketua": True}],
                           _IDENT, "Nusantara, 1 Januari 2026")
        sel = d.tables[0].rows[1].cells
        assert "Anggota," in _isi_sel(sel[0])[0]
        assert "Ketua Tim," in _isi_sel(sel[1])[0]

    def test_tanpa_penanda_ketua_yang_pertama_dipakai(self):
        d = dx.doc_baru()
        dx.signature_block(d, [{"nama": "A"}, {"nama": "B"}], _IDENT,
                           "Nusantara, 1 Januari 2026")
        sel = d.tables[0].rows[1].cells
        assert _isi_sel(sel[0])[0] == "Ketua Tim,"
        assert _isi_sel(sel[1])[0] == "Anggota,"

    def test_nip_asn_dicetak(self):
        d = dx.doc_baru()
        dx.signature_block(d, [{"nama": "Budi", "nip": _NIP}], _IDENT,
                           "Nusantara, 1 Januari 2026")
        assert f"NIP. {_NIP}" in _isi_sel(d.tables[0].rows[1].cells[0])

    def test_nik_tidak_pernah_dicetak(self):
        """PRIVASI. NIK adalah identitas kependudukan — dokumen ini beredar
        keluar satker. Bocor di sini tak bisa ditarik kembali."""
        d = dx.doc_baru()
        dx.signature_block(d, [{"nama": "Sari", "nip": _NIK}], _IDENT,
                           "Nusantara, 1 Januari 2026")
        isi = _isi_sel(d.tables[0].rows[1].cells[0])
        assert "Sari" in isi
        assert not any("NIP" in b or _NIK in b for b in isi)
        assert _NIK not in _xml_dokumen(d)

    def test_non_asn_tidak_dicetak_nomornya(self):
        d = dx.doc_baru()
        dx.signature_block(d, [{"nama": "Satpam Joko", "nip": _NIP}], _IDENT,
                           "Nusantara, 1 Januari 2026",
                           status_by_nip={_NIP: "non_asn"})
        isi = _isi_sel(d.tables[0].rows[1].cells[0])
        assert "Satpam Joko" in isi
        assert not any(_NIP in b for b in isi)

    def test_nip_kosong_dapat_garis_titik_BERLABEL_NETRAL(self):
        """Garis kosong diisi tangan setelah dicetak, jadi tak ada satu pun
        jalur kode yang bisa mendeteksi jenis nomornya. Label "NIP." di situ
        menebak — dan menebak salah untuk Non-ASN (NIK) maupun TNI/POLRI
        (NRP). Labelnya karena itu netral, benar untuk ketiganya."""
        from pegawai_utils import PLACEHOLDER_IDENTITAS
        d = dx.doc_baru()
        dx.signature_block(d, [{"nama": "Budi", "nip": ""}], _IDENT,
                           "Nusantara, 1 Januari 2026")
        isi = _isi_sel(d.tables[0].rows[1].cells[0])
        assert any(b.startswith(PLACEHOLDER_IDENTITAS[:12]) for b in isi), isi
        assert all(not b.startswith("NIP. .") for b in isi), isi

    def test_anggota_berupa_string_diterima(self):
        """Beberapa generator lama masih mengoper daftar nama polos."""
        d = dx.doc_baru()
        dx.signature_block(d, ["Budi", "Sari"], _IDENT, "Nusantara, 1 Januari 2026")
        sel = d.tables[0].rows[1].cells
        assert "Budi" in _isi_sel(sel[0]) and "Sari" in _isi_sel(sel[1])

    def test_tim_kosong_tetap_menyediakan_ruang_ttd(self):
        d = dx.doc_baru()
        dx.signature_block(d, [], _IDENT, "Nusantara, 1 Januari 2026")
        assert "________________________" in _isi_sel(d.tables[0].rows[1].cells[0])

    def test_tempat_tanggal_hanya_di_baris_pertama(self):
        d = dx.doc_baru()
        dx.signature_block(d, [{"nama": f"P{i}"} for i in range(4)], _IDENT,
                           "Nusantara, 1 Januari 2026")
        assert d.tables[0].rows[0].cells[1].text == "Nusantara, 1 Januari 2026"
        assert d.tables[1].rows[0].cells[1].text == ""

    def test_label_tim_bisa_diganti(self):
        d = dx.doc_baru()
        dx.signature_block(d, [{"nama": "A"}], _IDENT, "Nusantara, 1 Januari 2026",
                           label_tim="Panitia")
        assert "Panitia:" in _teks_paragraf(d)

    def test_saksi_ikut_bertanda_tangan(self):
        d = dx.doc_baru()
        dx.signature_block(d, [{"nama": "A"}], _IDENT, "Nusantara, 1 Januari 2026",
                           saksi=[{"nama": "Andi", "jabatan": "Pengelola BMN", "nip": _NIP2}])
        semua = [b for sel in _sel_ttd(d) for b in sel]
        assert "Saksi," in semua and "Andi" in semua
        assert "Pengelola BMN" in semua

    def test_kpb_mengetahui_di_kolom_tengah(self):
        d = dx.doc_baru()
        dx.signature_block(d, [{"nama": "A"}], _IDENT, "Nusantara, 1 Januari 2026")
        kpb = d.tables[-1]
        assert len(kpb.columns) == 3
        assert kpb.rows[0].cells[0].text == "" and kpb.rows[0].cells[2].text == ""
        isi = _isi_sel(kpb.rows[0].cells[1])
        assert isi[0] == "Mengetahui,"
        assert "Kuasa Pengguna Barang," in isi
        assert "Dr. Sartono" in isi
        assert f"NIP. {_IDENT['kasatker_nip']}" in isi

    def test_header_mengetahui_bisa_diganti(self):
        d = dx.doc_baru()
        dx.signature_block(d, [{"nama": "A"}], _IDENT, "Nusantara, 1 Januari 2026",
                           header_mengetahui="Menyetujui,")
        assert _isi_sel(d.tables[-1].rows[0].cells[1])[0] == "Menyetujui,"

    def test_kpb_tanpa_nip_tidak_mencetak_nomor_sampah(self):
        d = dx.doc_baru()
        dx.signature_block(d, [{"nama": "A"}], {"kasatker_nama": "Sartono"},
                           "Nusantara, 1 Januari 2026")
        from pegawai_utils import PLACEHOLDER_IDENTITAS
        teks = " ".join(_isi_sel(d.tables[-1].rows[0].cells[1]))
        assert PLACEHOLDER_IDENTITAS[:12] in teks, teks

    def test_ruang_tanda_tangan_tetap_lega(self):
        """Empat baris kosong = area tanda tangan basah. Menghapusnya membuat
        nama menempel pada jabatan dan tak ada tempat membubuhkan tanda
        tangan."""
        d = dx.doc_baru()
        dx.signature_block(d, [{"nama": "Budi", "nip": _NIP}], _IDENT,
                           "Nusantara, 1 Januari 2026")
        baris = _baris_sel(d.tables[0].rows[1].cells[0])
        kosong = sum(1 for b in baris if not b.strip())
        assert kosong >= 4

    def test_dokumen_akhir_tetap_sah(self):
        d = dx.doc_baru()
        dx.signature_block(d, [{"nama": "A", "nip": _NIP}, "B"], _IDENT,
                           "Nusantara, 1 Januari 2026",
                           saksi=[{"nama": "C"}])
        assert len(_buka_ulang(d).tables) == len(d.tables)


class TestTandaTanganTunggal:
    def test_di_kolom_kanan(self):
        d = dx.doc_baru()
        dx.signature_single(d, nama="Sartono", nip=_NIP)
        sel = d.tables[0].rows[0].cells
        assert sel[0].text == ""
        assert "Sartono" in _isi_sel(sel[1])

    def test_header_bawaan_pernyataan(self):
        d = dx.doc_baru()
        dx.signature_single(d, nama="Sartono")
        assert _isi_sel(d.tables[0].rows[0].cells[1])[0] == "Yang membuat pernyataan,"

    def test_jabatan_sebagai_peran_di_atas_ruang_ttd(self):
        d = dx.doc_baru()
        dx.signature_single(d, nama="Sartono", jabatan="Kuasa Pengguna Barang")
        isi = _isi_sel(d.tables[0].rows[0].cells[1])
        assert isi.index("Kuasa Pengguna Barang") < isi.index("Sartono")

    def test_jabatan_bawah_dicetak_setelah_nama(self):
        d = dx.doc_baru()
        dx.signature_single(d, nama="Sartono", jabatan="Kuasa Pengguna Barang",
                            jabatan_bawah=True, nip=_NIP)
        isi = _isi_sel(d.tables[0].rows[0].cells[1])
        assert isi.index("Sartono") < isi.index("Kuasa Pengguna Barang")
        assert isi.index("Kuasa Pengguna Barang") < isi.index(f"NIP. {_NIP}")

    def test_pre_lines_di_atas_header(self):
        d = dx.doc_baru()
        dx.signature_single(d, nama="Sartono", pre_lines=["Nusantara, 1 Januari 2026"])
        isi = _isi_sel(d.tables[0].rows[0].cells[1])
        assert isi[0] == "Nusantara, 1 Januari 2026"
        assert isi[1] == "Yang membuat pernyataan,"

    def test_nip_kosong_memakai_label_NETRAL(self):
        from pegawai_utils import PLACEHOLDER_IDENTITAS
        d = dx.doc_baru()
        dx.signature_single(d, nama="Sartono", nip="")
        teks = " ".join(_isi_sel(d.tables[0].rows[0].cells[1]))
        assert PLACEHOLDER_IDENTITAS[:12] in teks, teks

    def test_privasi_nik_juga_berlaku_di_ttd_tunggal(self):
        d = dx.doc_baru()
        dx.signature_single(d, nama="Sari", nip=_NIK)
        xml = _xml_dokumen(d)
        assert "Sari" in xml and _NIK not in xml

    def test_privasi_non_asn_di_ttd_tunggal(self):
        d = dx.doc_baru()
        dx.signature_single(d, nama="Joko", nip=_NIP, status="non_asn")
        xml = _xml_dokumen(d)
        assert "Joko" in xml and _NIP not in xml

    def test_lebar_dua_kolom_seimbang(self):
        d = dx.doc_baru()
        dx.signature_single(d, nama="Sartono")
        sel = d.tables[0].rows[0].cells
        assert sel[0].width == sel[1].width


class TestIdentityBlock:
    def test_intro_lalu_tabel(self):
        d = dx.doc_baru()
        dx.identity_block(d, [("Nama", "Budi"), ("NIP", _NIP)])
        assert _teks_paragraf(d)[0] == "Yang bertanda tangan di bawah ini:"
        assert [c.text for c in d.tables[0].rows[0].cells] == ["Nama", ":", "Budi"]

    def test_intro_bisa_dimatikan(self):
        d = dx.doc_baru()
        dx.identity_block(d, [("Nama", "Budi")], intro="")
        assert _teks_paragraf(d) == []
        assert len(d.tables) == 1


class TestTembusan:
    def test_daftar_kosong_tidak_menulis_apa_pun(self):
        d = dx.doc_baru()
        dx.tembusan(d, [])
        assert _teks_paragraf(d) == []

    def test_none_tidak_menulis_apa_pun(self):
        d = dx.doc_baru()
        dx.tembusan(d, None)
        assert _teks_paragraf(d) == []

    def test_hanya_spasi_dianggap_kosong(self):
        """Header 'Tembusan:' yang berdiri sendiri tanpa isi tampak cacat pada
        surat resmi."""
        d = dx.doc_baru()
        dx.tembusan(d, ["  ", "\t", ""])
        assert _teks_paragraf(d) == []

    def test_penomoran_urut_dari_satu(self):
        d = dx.doc_baru()
        dx.tembusan(d, ["Inspektur", "  Kepala Biro Umum  ", "Arsip"])
        assert _teks_paragraf(d) == ["Tembusan:", "1. Inspektur",
                                     "2. Kepala Biro Umum", "3. Arsip"]

    def test_baris_kosong_tidak_menggeser_nomor(self):
        d = dx.doc_baru()
        dx.tembusan(d, ["A", "   ", "B"])
        assert _teks_paragraf(d)[1:] == ["1. A", "2. B"]


# ---------------------------------------------------------------------------
# Rakitan penuh — seperti yang dilakukan generator laporan sesungguhnya
# ---------------------------------------------------------------------------

def test_dokumen_lengkap_terbentuk_dan_terbaca_ulang():
    d = dx.doc_baru(landscape=False)
    dx.kop_surat(d, _KOP)
    dx.title_block(d, "BERITA ACARA\nHASIL INVENTARISASI", nomor="BA-1/OIKN/2026")
    dx.section(d, "DASAR", "I")
    dx.para(d, "Peraturan Menteri Keuangan Nomor 181/PMK.06/2016.")
    dx.section(d, "HASIL", "II")
    dx.data_table(d, ["No", "Kode Barang", "NUP", "Nilai"],
                  [[1, "3.05.01.01.001", 1, 1500000],
                   [2, "3.05.01.01.002", 2, None]], align_right={3})
    dx.identity_block(d, [("Nama", "Budi"), ("NIP", _NIP)])
    dx.signature_block(d, [{"nama": "Budi", "nip": _NIP, "is_ketua": True},
                           {"nama": "Sari", "nip": _NIK}],
                       _IDENT, "Nusantara, 1 Januari 2026",
                       saksi=[{"nama": "Andi", "nip": _NIP2}])
    dx.tembusan(d, ["Inspektur"])
    dx.page_footer(d, "Berita Acara Hasil Inventarisasi")

    b = dx.to_bytes(d)
    assert b[:2] == b"PK"
    ulang = Document(io.BytesIO(b))
    assert len(ulang.tables) == len(d.tables)
    assert "OTORITA IBU KOTA NUSANTARA" in _teks_paragraf(ulang)
    # Privasi bertahan sampai berkas jadi (dibaca dari XML di dalam ZIP).
    xml = _xml_dokumen(d)
    assert _NIP in xml, "NIP ASN memang harus tercetak"
    assert _NIK not in xml, "NIK tak boleh ikut sampai berkas akhir"


@pytest.mark.parametrize("landscape", [False, True])
def test_orientasi_apa_pun_menghasilkan_berkas_sah(landscape):
    d = dx.doc_baru(landscape=landscape)
    dx.kop_surat(d, _KOP)
    dx.data_table(d, ["A", "B"], [[1, 2]])
    dx.page_footer(d, "x")
    assert Document(io.BytesIO(dx.to_bytes(d))).tables
